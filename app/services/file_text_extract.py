"""Extract plain text from uploaded field-research files (txt/md/csv/json/pdf/docx)."""

from __future__ import annotations

import io
from pathlib import Path

_MAX_CHARS_PER_FILE = 100_000

_TEXT_EXT = frozenset({".txt", ".md", ".csv", ".json", ".log", ".tsv"})


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """Return UTF-8 text; raises ValueError for unsupported types."""
    if len(data) > 8 * 1024 * 1024:
        raise ValueError("Файл больше 8 МБ")

    ext = Path(filename).suffix.lower()
    if not ext and data[:4] == b"%PDF":
        ext = ".pdf"

    if ext in _TEXT_EXT:
        text = data.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as e:  # pragma: no cover
            raise ValueError("PDF: установите зависимость pypdf") from e
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts)
    elif ext == ".docx":
        try:
            import docx
        except ImportError as e:  # pragma: no cover
            raise ValueError("Word: установите зависимость python-docx") from e
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        text = "\n".join(parts)
    else:
        raise ValueError(
            f"Формат «{ext or '?'}» не поддержан. Используйте .txt, .md, .csv, .json, .pdf или .docx"
        )

    text = text.strip()
    if len(text) > _MAX_CHARS_PER_FILE:
        text = text[:_MAX_CHARS_PER_FILE] + "\n\n[…фрагмент обрезан]"
    return text
